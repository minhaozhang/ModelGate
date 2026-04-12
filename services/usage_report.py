import asyncio
import csv
import json
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as QN
from docx.shared import Pt, RGBColor
from sqlalchemy import and_, func, literal_column, select

from core.config import api_keys_cache, providers_cache, proxy_logger as logger
from core.database import AnalysisRecord, ApiKey, RequestLogRead as RequestLog, async_session_maker
from services.analysis_store import (
    ANALYSIS_STATUS_FAILED,
    ANALYSIS_STATUS_PENDING,
    ANALYSIS_STATUS_RETRYING,
    ANALYSIS_STATUS_RUNNING,
    ANALYSIS_STATUS_SUCCESS,
    get_analysis_record_by_id,
    list_analysis_artifacts,
    list_analysis_subtasks,
    replace_analysis_subtasks,
    set_analysis_subtask_status,
    start_analysis_task,
    upsert_analysis_artifact,
    upsert_analysis_record,
)
from services.proxy import call_internal_model_via_proxy

ANALYSIS_TYPE_USAGE_REPORT = "usage_report"
REPORT_ROOT = Path("reports") / "usage"
USAGE_REPORT_TEMPLATE = {
    "id": "usage-report-pro",
    "version": "1.0",
    "name": "Usage Report Pro",
    "description": "Export reusable CSV/Markdown artifacts first, then assemble the final DOCX report.",
    "outputs": [
        "key_summary.csv",
        "daily_breakdown.csv",
        "model_breakdown.csv",
        "status_breakdown.csv",
        "summary.md",
        "overview.md",
        "awards.md",
        "final_report.docx",
    ],
}
USAGE_REPORT_STEPS = [
    {"key": "task_created", "label": "Task created", "max_attempts": 1},
    {"key": "querying_usage_data", "label": "Querying usage data", "max_attempts": 2},
    {"key": "exporting_key_summary_csv", "label": "Exporting key summary CSV", "max_attempts": 2},
    {"key": "exporting_daily_breakdown_csv", "label": "Exporting daily breakdown CSV", "max_attempts": 2},
    {"key": "exporting_model_breakdown_csv", "label": "Exporting model breakdown CSV", "max_attempts": 2},
    {"key": "exporting_status_breakdown_csv", "label": "Exporting status breakdown CSV", "max_attempts": 2},
    {"key": "exporting_summary_md", "label": "Exporting summary markdown", "max_attempts": 2},
    {"key": "calling_model_overview", "label": "Calling model for overview", "max_attempts": 2},
    {"key": "calling_model_awards", "label": "Calling model for awards", "max_attempts": 2},
    {"key": "exporting_ai_markdown", "label": "Exporting AI markdown", "max_attempts": 2},
    {"key": "assembling_docx", "label": "Assembling DOCX report", "max_attempts": 2},
]
_TOKEN_EXPR = func.coalesce(
    RequestLog.tokens["total_tokens"].as_integer(),
    RequestLog.tokens["estimated"].as_integer(),
    literal_column("0"),
)


def get_usage_report_template() -> dict:
    return {
        **USAGE_REPORT_TEMPLATE,
        "steps": [
            {
                "key": step["key"],
                "label": step["label"],
                "max_attempts": step["max_attempts"],
            }
            for step in USAGE_REPORT_STEPS
        ],
    }


def _task_signature(
    start_date: str,
    end_date: str,
    exclude_api_key_ids: list[int] | None = None,
    template_id: str = USAGE_REPORT_TEMPLATE["id"],
    template_version: str = USAGE_REPORT_TEMPLATE["version"],
) -> str:
    exclude_ids = sorted(set(exclude_api_key_ids or []))
    return json.dumps(
        {
            "start_date": start_date,
            "end_date": end_date,
            "exclude_api_key_ids": exclude_ids,
            "template_id": template_id,
            "template_version": template_version,
        },
        ensure_ascii=False,
        sort_keys=True,
    )


def _build_scope_key(start_date: str, end_date: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    return f"{start_date}:{end_date}:{timestamp}"


def _task_dir(task_id: int) -> Path:
    path = REPORT_ROOT / str(task_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _normalize_output_path(path: Path) -> str:
    return str(path).replace("\\", "/")


def _ensure_file_artifact(path_text: str) -> dict:
    path = Path(path_text)
    if not path.is_absolute():
        path = Path.cwd() / path
    if not path.exists() or path.stat().st_size <= 0:
        raise ValueError(f"Expected artifact file at {path} to exist and be non-empty")
    return {
        "path": _normalize_output_path(path),
        "size_bytes": path.stat().st_size,
    }


def _variance(values: list[int]) -> float:
    if not values:
        return 0.0
    avg = sum(values) / len(values)
    return sum((v - avg) ** 2 for v in values) / len(values)


async def _find_running_usage_report(signature: str) -> AnalysisRecord | None:
    async with async_session_maker() as session:
        result = await session.execute(
            select(AnalysisRecord)
            .where(AnalysisRecord.analysis_type == ANALYSIS_TYPE_USAGE_REPORT)
            .order_by(AnalysisRecord.created_at.desc())
        )
        for record in result.scalars().all():
            params = record.params_json or {}
            if params.get("signature") != signature:
                continue
            if record.status in {ANALYSIS_STATUS_PENDING, ANALYSIS_STATUS_RUNNING, ANALYSIS_STATUS_RETRYING}:
                return record
    return None


async def _create_usage_report_record(
    start_date: str,
    end_date: str,
    exclude_api_key_ids: list[int] | None = None,
) -> AnalysisRecord:
    exclude_ids = sorted(set(exclude_api_key_ids or []))
    signature = _task_signature(start_date, end_date, exclude_ids)
    scope_key = _build_scope_key(start_date, end_date)
    record = await upsert_analysis_record(
        ANALYSIS_TYPE_USAGE_REPORT,
        scope_key,
        status=ANALYSIS_STATUS_PENDING,
        progress="task_created",
        template_id=USAGE_REPORT_TEMPLATE["id"],
        template_version=USAGE_REPORT_TEMPLATE["version"],
        params_json={
            "start_date": start_date,
            "end_date": end_date,
            "exclude_api_key_ids": exclude_ids,
            "template": get_usage_report_template(),
            "signature": signature,
        },
        error="",
        content="",
        model_used="",
    )
    await replace_analysis_subtasks(record.id, USAGE_REPORT_STEPS)
    await set_analysis_subtask_status(
        record.id,
        "task_created",
        status=ANALYSIS_STATUS_SUCCESS,
        output={"created_at": datetime.now().isoformat()},
        increment_attempt=True,
    )
    return record


async def _run_subtask(
    record: AnalysisRecord,
    step_key: str,
    worker,
    validator=None,
):
    step = next(step for step in USAGE_REPORT_STEPS if step["key"] == step_key)
    max_attempts = int(step.get("max_attempts", 1) or 1)

    for attempt in range(1, max_attempts + 1):
        await upsert_analysis_record(
            record.analysis_type,
            record.scope_key,
            status=ANALYSIS_STATUS_RUNNING,
            progress=step_key,
            template_id=record.template_id,
            template_version=record.template_version,
            params_json=record.params_json,
        )
        await set_analysis_subtask_status(
            record.id,
            step_key,
            status=ANALYSIS_STATUS_RUNNING,
            increment_attempt=True,
            error="",
        )
        try:
            result = await worker()
            if validator:
                validator(result)
            await set_analysis_subtask_status(
                record.id,
                step_key,
                status=ANALYSIS_STATUS_SUCCESS,
                output=result if isinstance(result, dict) else {"value": result},
                error="",
            )
            return result
        except Exception as exc:
            is_last_attempt = attempt >= max_attempts
            await set_analysis_subtask_status(
                record.id,
                step_key,
                status=ANALYSIS_STATUS_FAILED if is_last_attempt else ANALYSIS_STATUS_RETRYING,
                output={"attempt": attempt, "max_attempts": max_attempts},
                error=str(exc),
            )
            if is_last_attempt:
                raise
            logger.warning(
                "[USAGE_REPORT] Retrying step %s for task %s after failure: %s",
                step_key,
                record.id,
                exc,
            )
            await asyncio.sleep(min(attempt, 3))


async def query_usage_stats(
    start_date: str,
    end_date: str,
    exclude_api_key_ids: list[int] | None = None,
) -> dict:
    start_dt = datetime.strptime(start_date, "%Y-%m-%d")
    end_dt = datetime.strptime(end_date, "%Y-%m-%d").replace(
        hour=23,
        minute=59,
        second=59,
    )

    async with async_session_maker() as session:
        base_filter = and_(
            RequestLog.created_at >= start_dt,
            RequestLog.created_at <= end_dt,
        )
        if exclude_api_key_ids:
            base_filter = and_(
                base_filter,
                RequestLog.api_key_id.notin_(exclude_api_key_ids),
            )

        keys_data: dict[int, dict] = {}

        agg_result = await session.execute(
            select(
                RequestLog.api_key_id,
                func.count().label("total_requests"),
                func.sum(_TOKEN_EXPR).label("total_tokens"),
                func.avg(RequestLog.latency_ms).label("avg_latency"),
            )
            .where(base_filter)
            .where(RequestLog.api_key_id.isnot(None))
            .group_by(RequestLog.api_key_id)
        )
        for row in agg_result:
            keys_data[row.api_key_id] = {
                "api_key_id": row.api_key_id,
                "total_requests": int(row.total_requests or 0),
                "total_tokens": int(row.total_tokens or 0),
                "avg_latency": round(row.avg_latency, 2) if row.avg_latency else 0,
                "status_distribution": {},
                "model_distribution": [],
                "hourly_distribution": {},
                "daily_distribution": {},
            }

        status_rows = []
        status_result = await session.execute(
            select(
                RequestLog.api_key_id,
                RequestLog.status,
                func.count().label("count"),
            )
            .where(base_filter)
            .where(RequestLog.api_key_id.isnot(None))
            .group_by(RequestLog.api_key_id, RequestLog.status)
        )
        for row in status_result:
            status_row = {
                "api_key_id": row.api_key_id,
                "status": row.status,
                "count": int(row.count or 0),
            }
            status_rows.append(status_row)
            if row.api_key_id in keys_data:
                keys_data[row.api_key_id]["status_distribution"][row.status] = status_row["count"]

        model_rows = []
        model_result = await session.execute(
            select(
                RequestLog.api_key_id,
                RequestLog.model,
                func.count().label("count"),
                func.sum(_TOKEN_EXPR).label("tokens"),
            )
            .where(base_filter)
            .where(RequestLog.api_key_id.isnot(None))
            .group_by(RequestLog.api_key_id, RequestLog.model)
        )
        raw_model_data: dict[int, list] = defaultdict(list)
        for row in model_result:
            model_row = {
                "api_key_id": row.api_key_id,
                "model": row.model,
                "count": int(row.count or 0),
                "tokens": int(row.tokens or 0),
            }
            model_rows.append(model_row)
            raw_model_data[row.api_key_id].append(model_row)
        for key_id, models in raw_model_data.items():
            if key_id in keys_data:
                keys_data[key_id]["model_distribution"] = sorted(
                    models,
                    key=lambda item: (item["count"], item["tokens"]),
                    reverse=True,
                )

        daily_rows = []
        daily_result = await session.execute(
            select(
                RequestLog.api_key_id,
                func.date(RequestLog.created_at).label("day"),
                func.count().label("count"),
            )
            .where(base_filter)
            .where(RequestLog.api_key_id.isnot(None))
            .group_by(RequestLog.api_key_id, func.date(RequestLog.created_at))
        )
        for row in daily_result:
            daily_row = {
                "api_key_id": row.api_key_id,
                "day": str(row.day),
                "count": int(row.count or 0),
            }
            daily_rows.append(daily_row)
            if row.api_key_id in keys_data:
                keys_data[row.api_key_id]["daily_distribution"][daily_row["day"]] = daily_row["count"]

        hourly_result = await session.execute(
            select(
                RequestLog.api_key_id,
                func.extract("hour", RequestLog.created_at).label("hour"),
                func.count().label("count"),
            )
            .where(base_filter)
            .where(RequestLog.api_key_id.isnot(None))
            .group_by(RequestLog.api_key_id, func.extract("hour", RequestLog.created_at))
        )
        for row in hourly_result:
            if row.api_key_id in keys_data:
                keys_data[row.api_key_id]["hourly_distribution"][str(int(row.hour))] = int(row.count or 0)

    api_key_names: dict[int, str] = {}
    for key_data in api_keys_cache.values():
        if isinstance(key_data, dict) and "id" in key_data:
            api_key_names[key_data["id"]] = key_data.get("name", f"Key-{key_data['id']}")

    missing_ids = [key_id for key_id in keys_data if key_id not in api_key_names]
    if missing_ids:
        async with async_session_maker() as session:
            db_result = await session.execute(
                select(ApiKey.id, ApiKey.name).where(ApiKey.id.in_(missing_ids))
            )
            for row in db_result:
                api_key_names[row.id] = row.name

    for key_id, data in keys_data.items():
        data["api_key_name"] = api_key_names.get(key_id, f"Key-{key_id}")

    sorted_keys = sorted(
        keys_data.values(),
        key=lambda item: (item["total_requests"], item["total_tokens"]),
        reverse=True,
    )
    total_requests = sum(item["total_requests"] for item in sorted_keys)
    total_tokens = sum(item["total_tokens"] for item in sorted_keys)
    overall_status = defaultdict(int)
    for row in status_rows:
        overall_status[row["status"]] += row["count"]

    return {
        "keys": sorted_keys,
        "start_date": start_date,
        "end_date": end_date,
        "generated_at": datetime.now().isoformat(),
        "totals": {
            "api_key_count": len(sorted_keys),
            "total_requests": total_requests,
            "total_tokens": total_tokens,
            "overall_status_distribution": dict(overall_status),
        },
        "rows": {
            "daily": daily_rows,
            "models": model_rows,
            "statuses": status_rows,
        },
    }


def _choose_analysis_model() -> tuple[str | None, str | None]:
    preferred_models = ("glm-5-turbo", "glm-5", "glm-4.6", "glm-4.7")
    preferred_providers = ("zhipu", "deepseek", "minimax", "ollama")

    for provider_name in preferred_providers:
        provider = providers_cache.get(provider_name)
        if not provider:
            continue
        for model_info in provider.get("models", []):
            actual_name = model_info.get("actual_model_name", "")
            if actual_name in preferred_models:
                return provider_name, actual_name

    for provider_name, provider in providers_cache.items():
        for model_info in provider.get("models", []):
            actual_name = model_info.get("actual_model_name", "")
            if actual_name:
                return provider_name, actual_name

    return None, None


def _build_summary_markdown(stats_data: dict) -> str:
    totals = stats_data.get("totals", {})
    lines = [
        "# Usage Report Summary",
        "",
        f"- Date Range: {stats_data.get('start_date')} to {stats_data.get('end_date')}",
        f"- API Keys: {totals.get('api_key_count', 0)}",
        f"- Total Requests: {totals.get('total_requests', 0)}",
        f"- Total Tokens: {totals.get('total_tokens', 0)}",
        "",
        "## Top API Keys",
    ]
    for index, key_info in enumerate(stats_data.get("keys", [])[:5], start=1):
        success_count = key_info.get("status_distribution", {}).get("success", 0)
        total_count = max(sum(key_info.get("status_distribution", {}).values()), 1)
        success_rate = (success_count / total_count) * 100
        top_model = (
            key_info.get("model_distribution", [{}])[0].get("model")
            if key_info.get("model_distribution")
            else "N/A"
        )
        lines.append(
            f"{index}. {key_info.get('api_key_name', 'Unknown')} | "
            f"requests={key_info.get('total_requests', 0)} | "
            f"tokens={key_info.get('total_tokens', 0)} | "
            f"success_rate={success_rate:.1f}% | top_model={top_model}"
        )
    if len(lines) == 7:
        lines.append("- No usage data")
    return "\n".join(lines).strip() + "\n"


def _build_awards_prompt(stats_data: dict) -> str:
    keys_summary = []
    for key_info in stats_data.get("keys", [])[:12]:
        hourly = key_info.get("hourly_distribution", {})
        daily = key_info.get("daily_distribution", {})
        daily_values = list(daily.values())
        keys_summary.append(
            {
                "name": key_info.get("api_key_name", "Unknown"),
                "total_requests": key_info.get("total_requests", 0),
                "total_tokens": key_info.get("total_tokens", 0),
                "avg_latency": key_info.get("avg_latency", 0),
                "model_count": len(key_info.get("model_distribution", [])),
                "night_requests": sum(hourly.get(str(h), 0) for h in range(0, 6)),
                "morning_requests": sum(hourly.get(str(h), 0) for h in range(6, 9)),
                "daily_variance": _variance(daily_values) if len(daily_values) > 1 else 0,
                "status_distribution": key_info.get("status_distribution", {}),
            }
        )

    if not keys_summary:
        return """你是一个专业但有趣的 API 使用报告分析师。请输出一份 markdown 奖项页。

要求：
1. 只输出 markdown。
2. 说明当前统计周期没有可用数据。
3. 结构至少包含：
   - # 趣味奖项
   - ## 总结
"""

    award_candidates = {
        "most_hardworking": max(keys_summary, key=lambda item: item["total_requests"]),
        "night_owl": max(keys_summary, key=lambda item: item["night_requests"]),
        "token_powerhouse": max(keys_summary, key=lambda item: item["total_tokens"]),
        "model_explorer": max(keys_summary, key=lambda item: item["model_count"]),
        "efficiency_king": min(
            [item for item in keys_summary if item["total_requests"] >= 3] or keys_summary,
            key=lambda item: item["avg_latency"] or 10**9,
        ),
    }

    return f"""你是一个专业但有趣的 API 使用报告分析师。请输出一份 markdown 奖项页。

要求：
1. 用中文输出。
2. 只输出 markdown，不要解释。
3. 结构包含：
   - # 趣味奖项
   - 每个奖项一个二级标题
   - 每个奖项包含“获得者”和“理由”
4. 奖项建议至少包含：
   - 最勤奋奖
   - 夜猫子奖
   - 输出狂人奖
   - 模型尝鲜者
   - 效率之王
5. 最后追加“## 总结”。
6. 只基于候选数据做判断，不要发明不存在的指标。

整体概览：
{json.dumps(stats_data.get("totals", {}), ensure_ascii=False, indent=2)}

奖项候选：
{json.dumps(award_candidates, ensure_ascii=False, indent=2)}
"""


def _build_overview_prompt(stats_data: dict) -> str:
    top_keys = []
    for item in stats_data.get("keys", [])[:6]:
        status_distribution = item.get("status_distribution", {})
        top_keys.append(
            {
                "name": item.get("api_key_name"),
                "requests": item.get("total_requests", 0),
                "tokens": item.get("total_tokens", 0),
                "avg_latency": item.get("avg_latency", 0),
                "success": status_distribution.get("success", 0),
                "errors": status_distribution.get("error", 0),
                "timeouts": status_distribution.get("timeout", 0),
                "rate_limited": status_distribution.get("rate_limited", 0),
                "top_model": (
                    item.get("model_distribution", [{}])[0].get("model", "N/A")
                    if item.get("model_distribution")
                    else "N/A"
                ),
            }
        )
    overview_payload = {
        "date_range": [stats_data.get("start_date"), stats_data.get("end_date")],
        "totals": stats_data.get("totals", {}),
        "top_keys": top_keys,
    }
    return f"""你是一个 API 平台运营分析师。请根据下面的数据输出一份中文 markdown 概览。

要求：
1. 只输出 markdown。
2. 结构包含：
   - # 报告概览
   - ## 关键发现
   - ## 风险信号
   - ## 建议动作
3. 结论要简洁、可执行，避免空话。

数据：
{json.dumps(overview_payload, ensure_ascii=False, indent=2)}
"""


async def _call_llm_markdown(stats_data: dict, prompt_kind: str) -> tuple[str | None, str | None]:
    provider_name, model_name = _choose_analysis_model()
    if not provider_name or not model_name:
        logger.warning("[USAGE_REPORT] No available model for AI markdown generation")
        return None, None

    if prompt_kind == "awards":
        prompt = _build_awards_prompt(stats_data)
        system_prompt = "你是一位擅长把数据转成清晰奖项页的分析师。"
    else:
        prompt = _build_overview_prompt(stats_data)
        system_prompt = "你是一位擅长生成高层运营报告概览的分析师。"

    body_json = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.4 if prompt_kind == "overview" else 0.7,
        "max_tokens": 2048,
    }

    try:
        result = await call_internal_model_via_proxy(
            requested_model=f"{provider_name}/{model_name}",
            body_json=body_json,
            purpose=f"usage-report-{prompt_kind}",
            timeout_seconds=120.0,
        )
        payload = result.get("payload")
        if not payload:
            logger.warning("[USAGE_REPORT] No payload for %s markdown: %s", prompt_kind, result.get("error"))
            return None, None

        message = ((payload.get("choices") or [{}])[0]).get("message") or {}
        content = message.get("content") or message.get("reasoning_content") or ""
        if isinstance(content, list):
            content = " ".join(
                item.get("text", "")
                for item in content
                if isinstance(item, dict)
            )
        return (content or None), f"{provider_name}/{model_name}"
    except Exception as exc:
        logger.warning("[USAGE_REPORT] Failed to generate %s markdown: %s", prompt_kind, exc)
        return None, None


def _write_csv(path: Path, headers: list[str], rows: list[dict]) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=headers)
        writer.writeheader()
        for row in rows:
            writer.writerow(row)
    return _normalize_output_path(path)


def _write_markdown(path: Path, markdown: str) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(markdown, encoding="utf-8")
    return _normalize_output_path(path)


def _append_markdown_to_doc(doc: Document, markdown_text: str) -> None:
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[:2].isdigit() and ". " in line:
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)


def _generate_docx(
    stats_data: dict,
    summary_markdown: str,
    overview_markdown: str,
    awards_markdown: str,
    output_path: str,
) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(QN("w:eastAsia"), "宋体")

    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ModelGate API Key 使用报告")
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"统计周期: {stats_data.get('start_date')} 至 {stats_data.get('end_date')}    生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _append_markdown_to_doc(doc, summary_markdown)
    _append_markdown_to_doc(doc, overview_markdown)
    _append_markdown_to_doc(doc, awards_markdown)

    doc.add_heading("附录：Top API Key 一览", level=1)
    keys = stats_data.get("keys", [])
    if keys:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["API Key", "请求数", "Token 用量", "平均延迟(ms)", "状态分布", "最常用模型"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for key_info in keys[:20]:
            row = table.add_row().cells
            status_text = ", ".join(
                f"{status}:{count}"
                for status, count in key_info.get("status_distribution", {}).items()
            ) or "N/A"
            top_model = key_info.get("model_distribution", [{}])[0].get("model", "N/A")
            row[0].text = key_info.get("api_key_name", "Unknown")
            row[1].text = str(key_info.get("total_requests", 0))
            row[2].text = str(key_info.get("total_tokens", 0))
            row[3].text = str(key_info.get("avg_latency", 0))
            row[4].text = status_text
            row[5].text = top_model
    else:
        doc.add_paragraph("暂无使用数据。")

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return _normalize_output_path(path)


def _serialize_subtask(record) -> dict:
    return {
        "id": record.id,
        "step_key": record.step_key,
        "step_label": record.step_label,
        "status": record.status,
        "sort_order": record.sort_order,
        "attempt_count": record.attempt_count or 0,
        "max_attempts": record.max_attempts or 1,
        "output": record.output or {},
        "error": record.error,
        "started_at": str(record.started_at) if record.started_at else None,
        "finished_at": str(record.finished_at) if record.finished_at else None,
    }


def _serialize_artifact(record) -> dict:
    return {
        "id": record.id,
        "subtask_id": record.subtask_id,
        "artifact_key": record.artifact_key,
        "artifact_type": record.artifact_type,
        "title": record.title,
        "path": record.path,
        "status": record.status,
        "meta": record.meta or {},
    }


async def _serialize_usage_report_record(record: AnalysisRecord) -> dict:
    subtasks = await list_analysis_subtasks(record.id)
    artifacts = await list_analysis_artifacts(record.id)
    result = {
        "id": record.id,
        "scope_key": record.scope_key,
        "status": record.status,
        "progress": record.progress,
        "model_used": record.model_used,
        "template_id": record.template_id,
        "template_version": record.template_version,
        "params": record.params_json or {},
        "error": record.error,
        "created_at": str(record.created_at) if record.created_at else None,
        "updated_at": str(record.updated_at) if record.updated_at else None,
        "subtasks": [_serialize_subtask(item) for item in subtasks],
        "artifacts": [_serialize_artifact(item) for item in artifacts],
    }
    if record.status == ANALYSIS_STATUS_SUCCESS and record.content:
        result["download_url"] = f"/admin/api/reports/usage/{record.id}/download"
    return result


async def generate_usage_report(task_id: int) -> None:
    record = await get_analysis_record_by_id(task_id)
    if not record:
        raise ValueError(f"Usage report task {task_id} not found")

    params = record.params_json or {}
    start_date = params.get("start_date")
    end_date = params.get("end_date")
    exclude_api_key_ids = params.get("exclude_api_key_ids") or []
    output_dir = _task_dir(record.id)
    task_state: dict[str, str | dict] = {}
    current_step = "task_created"

    try:
        current_step = "querying_usage_data"
        stats_data = await _run_subtask(
            record,
            "querying_usage_data",
            lambda: query_usage_stats(start_date, end_date, exclude_api_key_ids),
            validator=lambda result: result if isinstance(result, dict) and "keys" in result else (_ for _ in ()).throw(ValueError("Missing usage stats")),
        )

        async def export_key_summary_csv():
            path = output_dir / "key_summary.csv"
            csv_path = _write_csv(
                path,
                ["api_key_id", "api_key_name", "total_requests", "total_tokens", "avg_latency"],
                [
                    {
                        "api_key_id": item.get("api_key_id"),
                        "api_key_name": item.get("api_key_name"),
                        "total_requests": item.get("total_requests", 0),
                        "total_tokens": item.get("total_tokens", 0),
                        "avg_latency": item.get("avg_latency", 0),
                    }
                    for item in stats_data.get("keys", [])
                ],
            )
            task_state["key_summary_csv"] = csv_path
            return _ensure_file_artifact(csv_path)

        current_step = "exporting_key_summary_csv"
        key_summary_artifact = await _run_subtask(
            record,
            "exporting_key_summary_csv",
            export_key_summary_csv,
            validator=lambda result: _ensure_file_artifact(result["path"]),
        )
        key_summary_subtask = next(
            (item for item in await list_analysis_subtasks(record.id) if item.step_key == "exporting_key_summary_csv"),
            None,
        )
        await upsert_analysis_artifact(
            record.id,
            "key_summary_csv",
            artifact_type="csv",
            title="API Key Summary CSV",
            path=key_summary_artifact["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=key_summary_artifact,
            subtask_id=key_summary_subtask.id if key_summary_subtask else None,
        )

        async def export_daily_breakdown_csv():
            path = output_dir / "daily_breakdown.csv"
            csv_path = _write_csv(
                path,
                ["api_key_id", "day", "count"],
                stats_data.get("rows", {}).get("daily", []),
            )
            task_state["daily_breakdown_csv"] = csv_path
            return _ensure_file_artifact(csv_path)

        current_step = "exporting_daily_breakdown_csv"
        daily_artifact = await _run_subtask(
            record,
            "exporting_daily_breakdown_csv",
            export_daily_breakdown_csv,
            validator=lambda result: _ensure_file_artifact(result["path"]),
        )
        daily_subtask = next(
            (item for item in await list_analysis_subtasks(record.id) if item.step_key == "exporting_daily_breakdown_csv"),
            None,
        )
        await upsert_analysis_artifact(
            record.id,
            "daily_breakdown_csv",
            artifact_type="csv",
            title="Daily Breakdown CSV",
            path=daily_artifact["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=daily_artifact,
            subtask_id=daily_subtask.id if daily_subtask else None,
        )

        async def export_model_breakdown_csv():
            path = output_dir / "model_breakdown.csv"
            csv_path = _write_csv(
                path,
                ["api_key_id", "model", "count", "tokens"],
                stats_data.get("rows", {}).get("models", []),
            )
            task_state["model_breakdown_csv"] = csv_path
            return _ensure_file_artifact(csv_path)

        current_step = "exporting_model_breakdown_csv"
        model_artifact = await _run_subtask(
            record,
            "exporting_model_breakdown_csv",
            export_model_breakdown_csv,
            validator=lambda result: _ensure_file_artifact(result["path"]),
        )
        model_subtask = next(
            (item for item in await list_analysis_subtasks(record.id) if item.step_key == "exporting_model_breakdown_csv"),
            None,
        )
        await upsert_analysis_artifact(
            record.id,
            "model_breakdown_csv",
            artifact_type="csv",
            title="Model Breakdown CSV",
            path=model_artifact["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=model_artifact,
            subtask_id=model_subtask.id if model_subtask else None,
        )

        async def export_status_breakdown_csv():
            path = output_dir / "status_breakdown.csv"
            csv_path = _write_csv(
                path,
                ["api_key_id", "status", "count"],
                stats_data.get("rows", {}).get("statuses", []),
            )
            task_state["status_breakdown_csv"] = csv_path
            return _ensure_file_artifact(csv_path)

        current_step = "exporting_status_breakdown_csv"
        status_artifact = await _run_subtask(
            record,
            "exporting_status_breakdown_csv",
            export_status_breakdown_csv,
            validator=lambda result: _ensure_file_artifact(result["path"]),
        )
        status_subtask = next(
            (item for item in await list_analysis_subtasks(record.id) if item.step_key == "exporting_status_breakdown_csv"),
            None,
        )
        await upsert_analysis_artifact(
            record.id,
            "status_breakdown_csv",
            artifact_type="csv",
            title="Status Breakdown CSV",
            path=status_artifact["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=status_artifact,
            subtask_id=status_subtask.id if status_subtask else None,
        )

        async def export_summary_markdown():
            markdown = _build_summary_markdown(stats_data)
            path = output_dir / "summary.md"
            markdown_path = _write_markdown(path, markdown)
            task_state["summary_md"] = markdown_path
            task_state["summary_markdown_text"] = markdown
            return _ensure_file_artifact(markdown_path)

        current_step = "exporting_summary_md"
        summary_artifact = await _run_subtask(
            record,
            "exporting_summary_md",
            export_summary_markdown,
            validator=lambda result: _ensure_file_artifact(result["path"]),
        )
        summary_subtask = next(
            (item for item in await list_analysis_subtasks(record.id) if item.step_key == "exporting_summary_md"),
            None,
        )
        await upsert_analysis_artifact(
            record.id,
            "summary_md",
            artifact_type="markdown",
            title="Summary Markdown",
            path=summary_artifact["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=summary_artifact,
            subtask_id=summary_subtask.id if summary_subtask else None,
        )

        model_used_values: list[str] = []

        async def call_model_overview():
            markdown, model_used = await _call_llm_markdown(stats_data, "overview")
            if not markdown:
                markdown = "# 报告概览\n\n- 当前没有可用的模型分析结果，以下内容为基础导出报告。\n"
            if model_used:
                model_used_values.append(model_used)
            task_state["overview_markdown_text"] = markdown
            return {"markdown": markdown, "model_used": model_used or ""}

        current_step = "calling_model_overview"
        await _run_subtask(
            record,
            "calling_model_overview",
            call_model_overview,
            validator=lambda result: result if result.get("markdown") else (_ for _ in ()).throw(ValueError("Overview markdown missing")),
        )

        async def call_model_awards():
            markdown, model_used = await _call_llm_markdown(stats_data, "awards")
            if not markdown:
                markdown = "# 趣味奖项\n\n## 总结\n\n当前没有可用的模型奖项分析结果。\n"
            if model_used:
                model_used_values.append(model_used)
            task_state["awards_markdown_text"] = markdown
            return {"markdown": markdown, "model_used": model_used or ""}

        current_step = "calling_model_awards"
        await _run_subtask(
            record,
            "calling_model_awards",
            call_model_awards,
            validator=lambda result: result if result.get("markdown") else (_ for _ in ()).throw(ValueError("Awards markdown missing")),
        )

        async def export_ai_markdown():
            overview_path = _write_markdown(output_dir / "overview.md", task_state.get("overview_markdown_text", ""))
            awards_path = _write_markdown(output_dir / "awards.md", task_state.get("awards_markdown_text", ""))
            task_state["overview_md"] = overview_path
            task_state["awards_md"] = awards_path
            return {
                "overview": _ensure_file_artifact(overview_path),
                "awards": _ensure_file_artifact(awards_path),
            }

        current_step = "exporting_ai_markdown"
        ai_artifacts = await _run_subtask(
            record,
            "exporting_ai_markdown",
            export_ai_markdown,
            validator=lambda result: (_ensure_file_artifact(result["overview"]["path"]), _ensure_file_artifact(result["awards"]["path"])),
        )
        ai_subtask = next(
            (item for item in await list_analysis_subtasks(record.id) if item.step_key == "exporting_ai_markdown"),
            None,
        )
        await upsert_analysis_artifact(
            record.id,
            "overview_md",
            artifact_type="markdown",
            title="Overview Markdown",
            path=ai_artifacts["overview"]["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=ai_artifacts["overview"],
            subtask_id=ai_subtask.id if ai_subtask else None,
        )
        await upsert_analysis_artifact(
            record.id,
            "awards_md",
            artifact_type="markdown",
            title="Awards Markdown",
            path=ai_artifacts["awards"]["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=ai_artifacts["awards"],
            subtask_id=ai_subtask.id if ai_subtask else None,
        )

        async def assemble_docx():
            filename = f"{start_date}_{end_date}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            docx_path = _generate_docx(
                stats_data,
                task_state.get("summary_markdown_text", ""),
                task_state.get("overview_markdown_text", ""),
                task_state.get("awards_markdown_text", ""),
                str(output_dir / filename),
            )
            task_state["final_docx"] = docx_path
            return _ensure_file_artifact(docx_path)

        current_step = "assembling_docx"
        docx_artifact = await _run_subtask(
            record,
            "assembling_docx",
            assemble_docx,
            validator=lambda result: _ensure_file_artifact(result["path"]),
        )
        docx_subtask = next(
            (item for item in await list_analysis_subtasks(record.id) if item.step_key == "assembling_docx"),
            None,
        )
        await upsert_analysis_artifact(
            record.id,
            "final_docx",
            artifact_type="docx",
            title="Final DOCX Report",
            path=docx_artifact["path"],
            status=ANALYSIS_STATUS_SUCCESS,
            meta=docx_artifact,
            subtask_id=docx_subtask.id if docx_subtask else None,
        )

        await upsert_analysis_record(
            record.analysis_type,
            record.scope_key,
            status=ANALYSIS_STATUS_SUCCESS,
            progress="completed",
            content=task_state["final_docx"],
            model_used=", ".join(sorted(set(model_used_values))) if model_used_values else "rule-based",
            template_id=record.template_id,
            template_version=record.template_version,
            params_json=record.params_json,
            error="",
        )
    except Exception as exc:
        logger.error("[USAGE_REPORT] Failed to generate report %s: %s", task_id, exc)
        await upsert_analysis_record(
            record.analysis_type,
            record.scope_key,
            status=ANALYSIS_STATUS_FAILED,
            progress=current_step,
            error=str(exc),
            template_id=record.template_id,
            template_version=record.template_version,
            params_json=record.params_json,
        )
        raise


async def start_usage_report(
    start_date: str,
    end_date: str,
    exclude_api_key_ids: list[int] | None = None,
) -> tuple[bool, dict]:
    exclude_ids = sorted(set(exclude_api_key_ids or []))
    signature = _task_signature(start_date, end_date, exclude_ids)
    existing = await _find_running_usage_report(signature)
    if existing:
        return False, await _serialize_usage_report_record(existing)

    record = await _create_usage_report_record(start_date, end_date, exclude_ids)
    started = start_analysis_task(
        ANALYSIS_TYPE_USAGE_REPORT,
        signature,
        lambda: generate_usage_report(record.id),
    )
    if not started:
        existing = await _find_running_usage_report(signature)
        if existing:
            return False, await _serialize_usage_report_record(existing)
    return True, await _serialize_usage_report_record(record)


async def get_usage_report_status(task_id: int) -> dict | None:
    record = await get_analysis_record_by_id(task_id)
    if not record or record.analysis_type != ANALYSIS_TYPE_USAGE_REPORT:
        return None
    return await _serialize_usage_report_record(record)


async def list_usage_reports(limit: int = 10) -> list[dict]:
    async with async_session_maker() as session:
        result = await session.execute(
            select(AnalysisRecord)
            .where(AnalysisRecord.analysis_type == ANALYSIS_TYPE_USAGE_REPORT)
            .order_by(AnalysisRecord.created_at.desc())
            .limit(limit)
        )
        records = result.scalars().all()

    return [await _serialize_usage_report_record(record) for record in records]
